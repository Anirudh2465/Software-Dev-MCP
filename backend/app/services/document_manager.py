import os
import uuid
import json
import io
from pathlib import Path
import chromadb
from chromadb.utils import embedding_functions

# Optional libraries for file processing
try:
    import pypdf
    import docx
    from PIL import Image
    import pytesseract
except ImportError as e:
    print(f"Warning: Missing document processing dependencies: {e}")

# Load env
from dotenv import load_dotenv
load_dotenv()

CHROMA_HOST = "localhost"
CHROMA_PORT = 8000

class DocumentManager:
    def __init__(self):
        self.client = None
        self.collection = None
        
        try:
            print("Connecting to ChromaDB for Documents...")
            self.client = chromadb.HttpClient(host=CHROMA_HOST, port=CHROMA_PORT)
            # Use 'documents' collection
            self.collection = self.client.get_or_create_collection(
                name="documents",
                metadata={"hnsw:space": "cosine"}
            )
            print("Connected to ChromaDB 'documents' collection.")
        except Exception as e:
            print(f"Warning: Could not connect to ChromaDB for documents: {e}")

    def ingest_file(self, file_path_str: str, chunk_size=500):
        """
        Reads a file, chunks the text, and stores it in ChromaDB.
        """
        if not self.collection:
            return "Error: Document DB not available."

        path = Path(file_path_str)
        if not path.exists():
            return f"Error: File not found: {file_path_str}"
        
        try:
            text = self._extract_text(path)
            if not text:
                return f"Error: Could not extract text from {path.name} (empty or unsupported format)."
            
            # Chunking
            # Simple word-based chunking
            words = text.split()
            chunks = []
            for i in range(0, len(words), chunk_size):
                chunk = " ".join(words[i:i + chunk_size])
                chunks.append(chunk)
            
            # Indexing
            ids = [str(uuid.uuid4()) for _ in chunks]
            metadatas = [{"filename": path.name, "path": str(path.resolve()), "chunk_index": i} for i in range(len(chunks))]
            
            self.collection.upsert(
                ids=ids,
                documents=chunks,
                metadatas=metadatas
            )
            
            return f"Successfully ingested '{path.name}'. Created {len(chunks)} chunks."
            
        except Exception as e:
            return f"Error ingesting file {path.name}: {e}"

    def search_documents(self, query: str, n_results=3):
        """
        Retrieves relevant document chunks for a query.
        """
        if not self.collection:
            return []
            
        try:
            results = self.collection.query(
                query_texts=[query],
                n_results=n_results
            )
            
            docs = []
            if results['ids'] and results['ids'][0]:
                 for i, doc_text in enumerate(results['documents'][0]):
                     meta = results['metadatas'][0][i]
                     docs.append(f"[File: {meta['filename']}]\n{doc_text}")
            
            return docs
        except Exception as e:
            print(f"Error searching documents: {e}")
            return []

    def ingest_directory(self, dir_path: str, recursive=True):
        """
        Recursively ingests all supported files in a directory.
        """
        path = Path(dir_path)
        if not path.exists() or not path.is_dir():
            return f"Error: Directory not found: {dir_path}"

        results = []
        # Walk through directory
        for root, dirs, files in os.walk(path):
            if not recursive and root != str(path):
                continue
                
            for file in files:
                file_path = Path(root) / file
                # Skip system/hidden files or large binaries if needed
                if file.startswith("."): # Skip hidden files
                    continue
                    
                res = self.ingest_file(str(file_path))
                results.append(f"{file}: {res}")

        return f"Ingested directory {dir_path}. Results:\n" + "\n".join(results)

    def _extract_text(self, path: Path):
        ext = path.suffix.lower()
        
        if ext == ".pdf":
            try:
                reader = pypdf.PdfReader(str(path))
                text = ""
                for i, page in enumerate(reader.pages):
                    page_text = page.extract_text()
                    if page_text and page_text.strip():
                        text += page_text + "\n"
                    else:
                        # Try OCR on images
                        try:
                            if hasattr(page, 'images'):
                                for image_file in page.images:
                                    image_data = image_file.data
                                    img = Image.open(io.BytesIO(image_data))
                                    ocr_text = pytesseract.image_to_string(img)
                                    text += ocr_text + "\n"
                        except Exception as ocr_err:
                            print(f"Warning: OCR failed for page {i}: {ocr_err}")

                if not text.strip():
                     return "[WARNING: PDF appears empty or unreadable even after OCR attempt.]"
                return text
            except Exception as e:
                print(f"PDF Error: {e}")
                return None
                
        elif ext == ".docx":
            try:
                doc = docx.Document(str(path))
                text_parts = []
                
                # Extract paragraphs
                for para in doc.paragraphs:
                    if para.text.strip():
                        text_parts.append(para.text)
                
                # Extract tables
                for table in doc.tables:
                    for row in table.rows:
                        row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                        if row_text:
                            text_parts.append(" | ".join(row_text))
                
                text = "\n".join(text_parts)
                
                if not text.strip():
                    return "[WARNING: Docx contains no selectable text. It might be empty or contain only images.]"
                    
                return text
            except Exception as e:
                print(f"Docx Error: {e}")
                return None
        
        elif ext in [".txt", ".md", ".py", ".js", ".json", ".html", ".css", ".csv"]:
            try:
                with open(path, "r", encoding="utf-8", errors="ignore") as f:
                    return f.read()
            except Exception as e:
                 print(f"Text Read Error: {e}")
                 return None
                 
        elif ext in [".jpg", ".jpeg", ".png", ".bmp"]:
            try:
                # Requires Tesseract OCR installed on system
                # If tesseract is not in PATH, this raises TesseractNotFoundError
                image = Image.open(path)
                text = pytesseract.image_to_string(image)
                return text
            except Exception as e:
                print(f"OCR Error: {e} (Ensure Tesseract-OCR is installed and in PATH)")
                return f"[Image OCR Failed: {e}]"
        
        else:
            return f"[Unsupported file type: {ext}]"

